#!/usr/bin/env python
"""Minimal Markdown -> .docx converter for the SLR status/methodology notes.

Handles the subset of Markdown used in those files: ATX headings (#, ##, ###),
horizontal rules (---), blockquotes (>), unordered (-) and ordered (1.) lists,
pipe tables, and inline **bold** / `code`. Not a general converter; scoped to
keep the .docx twin of a hand/AI-authored .md in sync.

Usage: python md_to_docx.py INPUT.md OUTPUT.docx
"""
import re
import sys

from docx import Document
from docx.shared import Pt


def add_inline(paragraph, text):
    """Render **bold** and `code` runs into a paragraph."""
    # Split on bold (**...**) and inline code (`...`), keeping delimiters.
    tokens = re.split(r"(\*\*.*?\*\*|`[^`]*`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(tok)


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def split_row(line):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def main(src, dst):
    with open(src, encoding="utf-8") as fh:
        lines = fh.readlines()

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        # Table: header line followed by a separator line.
        if "|" in raw and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(raw)
            i += 2
            body = []
            while i < n and "|" in lines[i] and lines[i].strip():
                body.append(split_row(lines[i].rstrip("\n")))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, txt in enumerate(header):
                add_inline(table.rows[0].cells[c].paragraphs[0], txt)
            for row in body:
                cells = table.add_row().cells
                for c, txt in enumerate(row[: len(header)]):
                    add_inline(cells[c].paragraphs[0], txt)
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()  # thin separator
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        i += 1

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
